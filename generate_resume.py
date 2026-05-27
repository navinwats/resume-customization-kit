#!/usr/bin/env python3
"""
generate_resume.py — Generate a formatted .docx resume from a JSON content file.

Uses the MASTER resume as a style template so output exactly matches formatting.

Usage:
    python3 generate_resume.py content.json "Navin Watumull Resume MM.DD.YYYY Company Role.docx"

JSON format (see resume_instructions.md for field descriptions):
{
  "headline": "Marketing Leader | Integrated GTM, Brand & Product Storytelling",
  "summary": "Two-sentence summary with **bold anchors** inline.",
  "experience": [
    {
      "company": "ADOBE",
      "title": "Sr GTM Strategy Manager, Express",
      "location": "Los Angeles, CA",
      "dates": "Aug 2025-Present",
      "bullets": [
        "**Topic Header**: Body text with **bold metric** inline."
      ]
    }
  ],
  "skills": [
    { "header": "Integrated Marketing", "items": "GTM Strategy, Launch Planning, ..." }
  ],
  "education": "Bachelor of Arts, Economics (Cum Laude honors) | Tufts University, Medford, MA"
}
"""

import sys
import json
import os
import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Constants ────────────────────────────────────────────────────────────────

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MASTER_PATH = os.path.join(SCRIPT_DIR,
    "Navin Watumull MASTER Resume 04.14.2026 Master PMM Resume.docx")
OUTPUT_DIR = SCRIPT_DIR

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Right-aligned tab stop position: Letter page (12240) minus margins (1152 each) = 9936 DXA
RIGHT_TAB_POS = "9936"


# ── Low-level XML helpers ─────────────────────────────────────────────────────

def get_or_create_pPr(para):
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para._element.insert(0, pPr)
    return pPr


def get_or_create_rPr(run):
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    return rPr


def set_run_sz(run, sz_val):
    """Set explicit font size (in half-points) on a run."""
    rPr = get_or_create_rPr(run)
    for tag in ('w:sz', 'w:szCs'):
        el = OxmlElement(tag)
        el.set(qn('w:val'), str(sz_val))
        rPr.append(el)


def add_tab_run(para, sz_val=24):
    """Append a <w:r><w:rPr><w:sz/></w:rPr><w:tab/></w:r> to a paragraph."""
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(sz_val))
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(sz_val))
    rPr.append(sz)
    rPr.append(szCs)
    r.append(rPr)
    tab = OxmlElement('w:tab')
    r.append(tab)
    para._p.append(r)


def add_line_break(para):
    """Append a <w:r><w:br/></w:r> to a paragraph (soft line break within bullet)."""
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    r.append(br)
    para._p.append(r)


# ── Bold-marker parser ────────────────────────────────────────────────────────

def parse_bold(text):
    """
    Parse **bold** markers. Returns list of (segment_text, is_bold) tuples.
    Unclosed markers treated as plain text.
    """
    parts = []
    current = ""
    in_bold = False
    i = 0
    while i < len(text):
        if text[i:i+2] == "**":
            if current:
                parts.append((current, in_bold))
                current = ""
            in_bold = not in_bold
            i += 2
        else:
            current += text[i]
            i += 1
    if current:
        parts.append((current, in_bold))
    if in_bold:
        flat = "".join(seg for seg, _ in parts)
        return [(flat, False)]
    return parts


def add_runs(para, text, base_bold=False, sz=None):
    """
    Add runs with **bold** markers parsed.
    IMPORTANT: never set run.bold = False — that writes <w:b w:val="0"/> which
    overrides the paragraph style's bold, breaking Section Title etc.
    """
    for seg, is_bold in parse_bold(text):
        run = para.add_run(seg)
        if base_bold or is_bold:
            run.bold = True
        if sz is not None:
            set_run_sz(run, sz)


# ── Paragraph builder ─────────────────────────────────────────────────────────

def add_para(doc_element, text, style_name, base_bold=False):
    """
    Append a paragraph with style-specific handling:

    Section Title  — explicit sz=24, bold on all runs (style doesn't define size)
    List Paragraph — injects <w:numPr> for bullet marker (numId=2)
                     also splits on ' – ' or ' - ' (standalone dash/en-dash)
                     and inserts soft line breaks instead of connector dashes
    All others     — plain add_runs
    """
    para = doc_element.add_paragraph(style=style_name)

    if style_name == "Section Title":
        add_runs(para, text, base_bold=True, sz=24)

    elif style_name == "List Paragraph":
        pPr = get_or_create_pPr(para)
        numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '2')
        numPr.append(ilvl)
        numPr.append(numId)
        pPr.insert(0, numPr)
        # Explicitly set justification so both Word and LibreOffice honour it.
        # Without this, LibreOffice ignores the style-level JUSTIFY and renders
        # bullets as ragged-right even though Word shows them correctly.
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)
        # Split on standalone en dash or hyphen connectors → soft line break
        segments = re.split(r' [–\-] ', text)
        for i, segment in enumerate(segments):
            add_runs(para, segment, base_bold=base_bold)
            if i < len(segments) - 1:
                add_line_break(para)

    else:
        add_runs(para, text, base_bold=base_bold)

    return para


# ── Header — minimal in-place update ─────────────────────────────────────────

def update_headline_in_place(doc, headline):
    """
    Update only the headline text in the existing MASTER header table.
    Header cell paragraphs:
      0: Name
      1: Headline + <br/> + Contact (Body A, centered)
      2: Links (Body A, centered)
    Only touches paragraph 1: removes runs before <w:br/> run, inserts new bold headline run.
    Everything else (name, contact, links, hyperlinks, sizes) is left untouched.
    """
    cell = doc.tables[0].rows[0].cells[0]
    hl_para = cell.paragraphs[1]
    p = hl_para._p

    # Find the <w:r> that contains <w:br/> — contact line starts here
    br_parent_r = None
    for child in list(p):
        if child.tag == qn('w:r'):
            if child.find(qn('w:br')) is not None:
                br_parent_r = child
                break

    # Remove all <w:r> elements before br_parent_r
    for child in list(p):
        if child is br_parent_r:
            break
        if child.tag == qn('w:r'):
            p.remove(child)

    # Insert new bold headline run immediately before br_parent_r
    new_r = OxmlElement('w:r')
    new_rPr = OxmlElement('w:rPr')
    b = OxmlElement('w:b')
    new_rPr.append(b)
    new_r.append(new_rPr)
    new_t = OxmlElement('w:t')
    new_t.text = headline
    new_r.append(new_t)

    if br_parent_r is not None:
        br_parent_r.addprevious(new_r)
    else:
        p.append(new_r)


# ── Section builders ──────────────────────────────────────────────────────────

def build_summary(doc, summary_text):
    add_para(doc, "Summary", "Section Title")
    add_para(doc, summary_text, "Body A")


def add_role_spacer(doc):
    """
    Add a blank 6pt Normal paragraph before a role line.
    Matches the hand-built pattern: blank para with sz=12 in pPr rPr.
    """
    para = doc.add_paragraph()  # Normal style
    pPr = get_or_create_pPr(para)
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '12')
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '12')
    rPr.append(sz)
    rPr.append(szCs)
    pPr.append(rPr)
    return para


def add_role_line(doc, company, title, location, dates):
    """
    Build a company/role line matching hand-built resume formatting:
      COMPANY (bold sz=24), , (normal sz=24), Title (italic sz=24),
      " | " + location (normal sz=24), one right-aligned tab, then dates.

    Right-aligned tab stop at RIGHT_TAB_POS forces dates flush to right margin
    regardless of text length.
    keepNext prevents orphaned role headers at page breaks.
    """
    para = doc.add_paragraph(style=doc.styles["Body A"])

    # Right-aligned tab stop + keepNext in paragraph properties
    pPr = get_or_create_pPr(para)

    keep_next = OxmlElement('w:keepNext')
    pPr.append(keep_next)

    tabs_el = OxmlElement('w:tabs')
    tab_el = OxmlElement('w:tab')
    tab_el.set(qn('w:val'), 'right')
    tab_el.set(qn('w:pos'), RIGHT_TAB_POS)
    tabs_el.append(tab_el)
    pPr.append(tabs_el)

    def sized_run(text, bold=False, italic=False):
        run = para.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        set_run_sz(run, 24)
        return run

    if company:
        sized_run(company, bold=True)
        sized_run(", ")
    sized_run(title, italic=True)
    sized_run(f" | {location}")
    # Single right-aligned tab jumps to the right margin
    add_tab_run(para, sz_val=24)
    # Dates — normal weight, sz=22 (inherits from BodyA style)
    para.add_run(dates)
    return para


def build_experience(doc, experience):
    add_para(doc, "professional experience", "Section Title")

    for i, role in enumerate(experience):
        add_role_spacer(doc)
        add_role_line(
            doc,
            role.get("company", ""),
            role.get("title", ""),
            role.get("location", ""),
            role.get("dates", ""),
        )
        for bullet in role.get("bullets", []):
            add_para(doc, bullet, "List Paragraph")


def build_skills(doc, skills):
    add_para(doc, "Skills", "Section Title")
    for skill in skills:
        add_para(doc, f"**{skill.get('header', '')}**: {skill.get('items', '')}", "Body A")


def build_education(doc, education):
    add_para(doc, "Education", "Section Title")
    add_para(doc, education, "Body A")


# ── Document settings ─────────────────────────────────────────────────────────

def disable_auto_hyphenation(doc):
    """
    Remove <w:autoHyphenation/> from document settings.
    The MASTER has this enabled, which causes Word to silently break words
    mid-syllable (Prod-uct, Strat-egy) on justified lines. Removing it
    prevents all automatic hyphenation.
    """
    settings = doc.settings.element
    for el in settings.findall(f"{{{W}}}autoHyphenation"):
        settings.remove(el)


# ── Body cleaner ──────────────────────────────────────────────────────────────

def clear_document_body(doc):
    """
    Remove all body paragraphs while preserving:
      - All <w:tbl> elements (header table must survive)
      - <w:sectPr> (page margins / section properties)
    """
    body = doc.element.body
    sect_pr = body.find(f"{{{W}}}sectPr")
    tables = set(body.findall(f"{{{W}}}tbl"))
    preserve = tables | {sect_pr}
    for child in list(body):
        if child not in preserve:
            body.remove(child)


# ── Font normalization ────────────────────────────────────────────────────────

def _replace_font(doc, old_font, new_font):
    """Replace every occurrence of old_font with new_font across the document XML.

    Used to swap Word-bundled Garamond for EB Garamond (a system-installed
    Garamond-family font accessible to both Word and LibreOffice), ensuring the
    DOCX and LibreOffice-generated PDF look identical.
    """
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    rFonts_tag = f"{{{W_NS}}}rFonts"
    font_attrs = [
        f"{{{W_NS}}}ascii",
        f"{{{W_NS}}}hAnsi",
        f"{{{W_NS}}}eastAsia",
        f"{{{W_NS}}}cs",
    ]
    for tree in (doc.element, doc.styles.element):
        for el in tree.iter(rFonts_tag):
            for attr in font_attrs:
                if el.get(attr) == old_font:
                    el.set(attr, new_font)


# ── PDF export ───────────────────────────────────────────────────────────────

def _export_pdf(docx_path, pdf_path):
    """
    Convert .docx to .pdf. Tries two approaches in order:
      1. Word via direct osascript — no file picker, no docx2pdf dialogs.
         Word renders justified text correctly; LibreOffice does not.
         macOS will prompt for Automation access the very first time only.
      2. LibreOffice headless — fallback if Word is unavailable.
      3. Warns if both fail.
    """
    import subprocess, shutil

    docx_abs = os.path.abspath(docx_path)
    pdf_abs  = os.path.abspath(pdf_path)

    # ── Option 1: Word via osascript (no dialogs, correct justification) ──
    script = f'''
tell application "Microsoft Word"
    open POSIX file "{docx_abs}"
    set theDoc to active document
    save as theDoc file name "{pdf_abs}" file format format PDF
    close theDoc saving no
end tell
'''
    result = subprocess.run(
        ['osascript', '-e', script],
        capture_output=True, text=True
    )
    if result.returncode == 0 and os.path.exists(pdf_abs):
        print(f"✓ PDF (Word):   {pdf_abs}")
        return

    # ── Option 2: LibreOffice headless (fallback) ──
    soffice = (
        shutil.which("soffice") or
        shutil.which("libreoffice") or
        "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    )
    if soffice and os.path.exists(soffice):
        import pathlib
        out_dir = str(pathlib.Path(pdf_path).parent)
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
            capture_output=True
        )
        lo_pdf = docx_path.replace(".docx", ".pdf")
        if os.path.exists(lo_pdf) and lo_pdf != pdf_path:
            os.rename(lo_pdf, pdf_path)
        if os.path.exists(pdf_path):
            print(f"✓ PDF (LibreOffice): {pdf_path}")
            return

    print("⚠ PDF export skipped (both Word and LibreOffice failed).")
    if result.stderr:
        print(f"  Word error: {result.stderr.strip()}")
    print("  → Make sure Microsoft Word is installed for best PDF output.")


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) < 3:
        print("Usage: python3 generate_resume.py content.json \"Output Filename.docx\"")
        sys.exit(1)

    content_path = sys.argv[1]
    output_name  = sys.argv[2]

    with open(content_path, "r", encoding="utf-8") as f:
        content = json.load(f)

    doc = Document(MASTER_PATH)

    # Kill auto-hyphenation inherited from MASTER settings
    disable_auto_hyphenation(doc)

    # Update only the headline text — leave name, contact, links untouched
    update_headline_in_place(doc, content.get("headline", ""))

    # Clear body paragraphs (preserves header table and sectPr)
    clear_document_body(doc)

    build_summary(doc, content.get("summary", ""))
    build_experience(doc, content.get("experience", []))
    build_skills(doc, content.get("skills", []))
    build_education(doc, content.get("education", ""))

    if not output_name.lower().endswith(".docx"):
        output_name += ".docx"
    output_path = os.path.join(OUTPUT_DIR, output_name)

    # Swap Word-bundled Garamond for EB Garamond (installed system font) so that
    # both the DOCX and the LibreOffice-generated PDF use the same font.
    # EB Garamond is a high-quality Garamond-family font accessible to both Word
    # and LibreOffice. Install from: fonts.google.com/specimen/EB+Garamond
    _replace_font(doc, "Garamond", "EB Garamond")

    doc.save(output_path)
    print(f"✓ Saved: {output_path}")

    # Auto-export PDF alongside the .docx
    pdf_path = output_path.replace(".docx", ".pdf")
    _export_pdf(output_path, pdf_path)


if __name__ == "__main__":
    main()
